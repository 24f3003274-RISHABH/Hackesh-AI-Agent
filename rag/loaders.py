"""
Document Loaders for Hackesh Local RAG Engine.
Supports PDF, TXT, Markdown, and DOCX formats without external API calls.
"""

import os
from typing import Dict, Any, Tuple
from .models import Document


class DocumentLoader:
    """Loads and extracts raw text + metadata from local files."""

    @staticmethod
    def load_from_text(content: str, filename: str = "document.txt", doc_type: str = "txt", metadata: Dict[str, Any] = None) -> Document:
        """Loads a document directly from raw string content."""
        ext = filename.split(".")[-1].lower() if "." in filename else doc_type
        return Document(
            filename=filename,
            doc_type=ext,
            content=content,
            metadata=metadata or {},
            total_pages=1
        )

    @staticmethod
    def load_file(file_path: str) -> Document:
        """Loads a file from disk based on extension."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        filename = os.path.basename(file_path)
        ext = filename.split(".")[-1].lower()

        if ext in ["txt", "md", "markdown"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return Document(
                filename=filename,
                doc_type="md" if "md" in ext else "txt",
                content=content,
                metadata={"file_path": file_path, "size_bytes": len(content)}
            )

        elif ext == "pdf":
            # Extract text from PDF
            content, total_pages = DocumentLoader._extract_pdf(file_path)
            return Document(
                filename=filename,
                doc_type="pdf",
                content=content,
                total_pages=total_pages,
                metadata={"file_path": file_path, "pages": total_pages}
            )

        elif ext in ["docx", "doc"]:
            content = DocumentLoader._extract_docx(file_path)
            return Document(
                filename=filename,
                doc_type="docx",
                content=content,
                metadata={"file_path": file_path}
            )

        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return Document(
                filename=filename,
                doc_type="txt",
                content=content,
                metadata={"file_path": file_path}
            )

    @staticmethod
    def _extract_pdf(file_path: str) -> Tuple[str, int]:
        """Extracts text page-by-page from a PDF."""
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages_text.append(f"--- [Page {idx + 1}] ---\n{text}")
            return "\n\n".join(pages_text), len(reader.pages)
        except Exception:
            # Fallback text extraction if pypdf is not available
            with open(file_path, "rb") as f:
                raw = f.read()
            # Simple text extraction filter
            clean_chars = "".join([chr(b) if 32 <= b <= 126 or b in (10, 13) else " " for b in raw])
            return clean_chars, 1

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extracts text from a DOCX file."""
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
